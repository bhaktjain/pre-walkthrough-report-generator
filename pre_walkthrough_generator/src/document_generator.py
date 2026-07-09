from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime
import logging
import requests
from io import BytesIO
from PIL import Image

logger = logging.getLogger(__name__)

# Brand heading color (#231f20) applied to every report heading.
HEADING_COLOR = RGBColor(0x23, 0x1F, 0x20)

# Network timeout (seconds) for downloading images / floor plans.
IMAGE_DOWNLOAD_TIMEOUT = 15

# Cap how many neighboring projects are listed in the report so a large
# neighborhood doesn't produce a giant table. The list is pre-sorted
# (same-building first, then by amount), so the most relevant are shown.
MAX_NEIGHBORING_PROJECTS = 15


def _to_number(value) -> Optional[float]:
    """Best-effort numeric coercion. Returns a float, or None if not numeric.

    Handles ints/floats and money-formatted strings like "$1,250" / "1250.50".
    booleans are intentionally rejected (True is not a price).
    """
    if isinstance(value, bool) or value is None:
        return None
    if isinstance(value, (int, float)):
        return float(value)
    if isinstance(value, str):
        cleaned = value.replace(',', '').replace('$', '').strip()
        if not cleaned:
            return None
        try:
            return float(cleaned)
        except ValueError:
            return None
    return None


def _leading_money(value) -> Optional[float]:
    """Extract a leading currency amount even when trailing prose follows, e.g.
    "$736,000 (sold Sep 25, 2025)" -> 736000.0. Returns None if there's no
    leading number. Used only for price display so a model-annotated price still
    renders as a clean dollar figure (and doesn't duplicate the sale date)."""
    n = _to_number(value)
    if n is not None:
        return n
    import re as _re
    m = _re.match(r'\s*\$?\s*(\d[\d,]*(?:\.\d+)?)', str(value or ''))
    if m:
        try:
            return float(m.group(1).replace(',', ''))
        except ValueError:
            return None
    return None


# Palette (matches the report preview): brand near-black + warm neutrals.
LABEL_FILL = 'F4F2EE'    # shaded label column
ZEBRA_FILL = 'FAF9F7'    # alternating value rows
HEADER_FILL = '231F20'   # dark header row (neighbors table)
BORDER_COLOR = 'DDDAD4'  # light table borders
RULE_COLOR = 'C9C3BA'    # subtle rule under section headings
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _shade_cell(cell, fill_hex: str) -> None:
    """Apply a solid background fill to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def _heading_rule(paragraph, color_hex: str, sz: str) -> None:
    """Add a bottom border (divider rule) under a heading paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pbdr.append(bottom)
    pPr.append(pbdr)


def _table_borders(table, color_hex: str, sz: str = '4') -> None:
    """Set light, uniform borders on a table (replaces the heavy black grid)."""
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), color_hex)
        borders.append(e)
    tblPr.append(borders)


class DocumentGenerator:
    def __init__(self):
        self.doc = Document()
        self._setup_document()

    def _setup_document(self):
        """Set up document styles and formatting"""
        # Custom heading styles (kept for backwards compatibility) use the brand color.
        for i in range(1, 4):
            style = self.doc.styles.add_style(f'CustomHeading{i}', WD_STYLE_TYPE.PARAGRAPH)
            font = style.font
            font.size = Pt(16 - (i * 2))
            font.bold = True
            font.color.rgb = HEADING_COLOR

        # Apply the brand heading color to the built-in heading styles that
        # add_heading() actually uses (Title + Heading 1-4).
        for style_name in ('Title', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4'):
            try:
                self.doc.styles[style_name].font.color.rgb = HEADING_COLOR
            except (KeyError, AttributeError):
                pass

        # Add header and footer
        self._add_header_footer()

    # ------------------------------------------------------------------ helpers
    def _heading(self, text: str, level: int = 1):
        """Add a heading and force the brand color on its runs.

        Setting the style color (above) covers most cases, but setting the run
        color directly guarantees #231f20 regardless of theme-color behavior.
        """
        heading = self.doc.add_heading(text, level)
        for run in heading.runs:
            run.font.color.rgb = HEADING_COLOR
        # Divider rules: a bold brand rule under the title, a subtle rule under
        # each section heading (matches the report preview's sectioned look).
        if level == 0:
            _heading_rule(heading, HEADER_FILL, '12')
        elif level == 1:
            _heading_rule(heading, RULE_COLOR, '6')
        return heading

    @staticmethod
    def _safe_int(value) -> Optional[int]:
        num = _to_number(value)
        if num is None:
            return None
        try:
            return int(num)
        except (ValueError, OverflowError):
            return None

    @staticmethod
    def _zillow_search_url(address) -> Optional[str]:
        """A Zillow listing/search URL for an address (no API needed; Zillow
        resolves the address query to the listing, which often has a floor plan)."""
        if not address:
            return None
        from urllib.parse import quote
        return "https://www.zillow.com/homes/" + quote(str(address).strip()) + "_rb/"

    @staticmethod
    def _is_meaningful(value) -> bool:
        """True if a value carries real content.

        Recurses into dicts/lists so template defaults like {"range": {"min": 0,
        "max": None}} or {"type": "", "preferences": []} are correctly treated as
        empty (this is what produced 'min: 0' / '$0 - TBD' / blank rows before).
        Zero numbers and blank strings are NOT meaningful.
        """
        if value is None:
            return False
        if isinstance(value, bool):
            return value
        if isinstance(value, (int, float)):
            return value != 0
        if isinstance(value, str):
            return bool(value.strip())
        if isinstance(value, dict):
            return any(DocumentGenerator._is_meaningful(v) for v in value.values())
        if isinstance(value, (list, tuple)):
            return any(DocumentGenerator._is_meaningful(v) for v in value)
        return True

    def _lines(self, value) -> list:
        """Normalize a value to a list of non-empty display strings."""
        if value in (None, '', [], {}):
            return []
        if isinstance(value, str):
            return [value]
        if isinstance(value, dict):
            return [self._stringify(value)]
        if isinstance(value, (list, tuple)):
            out = []
            for v in value:
                if v in (None, '', [], {}):
                    continue
                out.append(self._stringify(v) if isinstance(v, (dict, list, tuple)) else str(v))
            return out
        return [str(value)]

    def _format_value(self, value, prefix="", suffix=""):
        """Format a value with optional prefix and suffix"""
        if value is None:
            return "Information not available"
        if prefix == "$":
            num = _to_number(value)
            if num is not None:
                rendered = int(num) if num == int(num) else num
                return f"{prefix}{rendered:,}{suffix}"
        return f"{prefix}{value}{suffix}"

    def _format_currency(self, value, prefix: str = "$") -> str:
        """Format a currency value, coercing numeric strings; 'TBD' if unknown."""
        num = _to_number(value)
        if num is None:
            return "TBD"
        if num == int(num):
            return f"{prefix}{int(num):,}"
        return f"{prefix}{num:,.2f}"

    def _add_header_footer(self):
        """Add header and footer images to the document"""
        try:
            header_path = "header.png"
            footer_path = "footer.png"

            if not Path(header_path).exists() or not Path(footer_path).exists():
                logger.warning(
                    "Header or footer image not found (header=%s, footer=%s)",
                    Path(header_path).exists(), Path(footer_path).exists(),
                )
                return

            section = self.doc.sections[0]

            # Header — always use a fresh run for the picture (never reuse runs[0],
            # which can corrupt an existing template run).
            header = section.header
            header_para = header.paragraphs[0]
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_para.add_run().add_picture(header_path, height=Inches(0.8))
            header_para.paragraph_format.space_after = Pt(12)

            # Footer
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.add_run().add_picture(footer_path, height=Inches(0.8))

            logger.info("Header and footer images added successfully")

        except Exception as e:
            logger.error("Error adding header/footer: %s", e)

    def _add_header(self, data: Dict[str, Any]):
        """Add document header with title and address"""
        title = self._heading('Pre-Walkthrough Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        address = data.get('property_address', '')
        if address:
            addr_para = self.doc.add_paragraph()
            addr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            addr_para.add_run(str(address)).bold = True

        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(datetime.now().strftime("%B %d, %Y"))

        self.doc.add_paragraph()

    def _add_executive_summary(self, data: Dict[str, Any]):
        """Add concise, sales-focused executive summary"""
        self._heading('Executive Summary', level=1)

        transcript = data.get('transcript_info', {}) or {}
        scope = transcript.get('renovation_scope', {}) or {}

        # 1. Project Goals (comprehensive project description)
        goals_parts = []

        kitchen = scope.get('kitchen', {}) or {}
        if kitchen.get('description'):
            goals_parts.append(f"Kitchen: {kitchen['description']}")

        bathrooms = scope.get('bathrooms', {}) or {}
        if bathrooms.get('plumbing_changes') or bathrooms.get('specific_requirements'):
            bath_desc = []
            if bathrooms.get('plumbing_changes'):
                bath_desc.append(str(bathrooms['plumbing_changes']))
            bath_desc.extend(self._lines(bathrooms.get('specific_requirements')))
            goals_parts.append(f"Bathroom: {', '.join(bath_desc)}")

        additional = scope.get('additional_work', {}) or {}
        if additional.get('rooms') or additional.get('structural_changes') or additional.get('systems_updates'):
            add_desc = []
            rooms = self._lines(additional.get('rooms'))
            structural = self._lines(additional.get('structural_changes'))
            systems = self._lines(additional.get('systems_updates'))
            if rooms:
                add_desc.append(f"Rooms: {', '.join(rooms)}")
            if structural:
                add_desc.append(f"Structural: {', '.join(structural)}")
            if systems:
                add_desc.append(f"Systems: {', '.join(systems)}")
            goals_parts.append(f"Additional work: {', '.join(add_desc)}")

        project_goals = '. '.join(goals_parts) if goals_parts else None

        # 2. Client Drivers
        drivers = []
        if additional.get('structural_changes'):
            drivers.append('layout upgrade')
        if additional.get('systems_updates'):
            drivers.append('systems modernisation')
        client_drivers = ', '.join(drivers) if drivers else None

        # 3. Key numbers — budget + timeline
        timeline = scope.get('timeline', {}) or {}
        duration = timeline.get('total_duration', 'TBD')

        estimated_costs = additional.get('estimated_costs') or {}
        per_sqft_cost = _to_number(estimated_costs.get('per_sqft_cost'))
        total_range = estimated_costs.get('total_estimated_range') or {}
        tr_min = _to_number(total_range.get('min'))
        tr_max = _to_number(total_range.get('max'))

        budget_phrase = None
        if per_sqft_cost is not None:
            if tr_min is not None and tr_max is not None:
                budget_phrase = f"Budget {self._format_currency(per_sqft_cost)}/sq ft (estimated {self._format_currency(tr_min)} – {self._format_currency(tr_max)})"
            elif tr_min is not None:
                budget_phrase = f"Budget {self._format_currency(per_sqft_cost)}/sq ft (estimated from {self._format_currency(tr_min)})"
            else:
                budget_phrase = f"Budget from {self._format_currency(per_sqft_cost)}/sq ft"
        elif tr_min is not None or tr_max is not None:
            lo = tr_min if tr_min is not None else tr_max
            hi = tr_max if tr_max is not None else tr_min
            budget_phrase = f"Budget {self._format_currency(lo)} – {self._format_currency(hi)}"
        else:
            # Sum up component costs as a fallback estimate.
            min_total, max_total = 0.0, 0.0
            have = False
            krange = (kitchen.get('estimated_cost') or {}).get('range') or {}
            kmin = _to_number(krange.get('min'))
            kmax = _to_number(krange.get('max'))
            if kmin is not None:
                min_total += kmin
                max_total += kmax if kmax is not None else kmin
                have = True
            bcount = self._safe_int(bathrooms.get('count'))
            bcost = _to_number(bathrooms.get('cost_per_bathroom'))
            if bcount and bcost:
                min_total += bcost * bcount
                max_total += bcost * bcount
                have = True
            if have:
                if max_total and max_total != min_total:
                    budget_phrase = f"Budget {self._format_currency(min_total)} – {self._format_currency(max_total)}"
                elif min_total:
                    budget_phrase = f"Budget from {self._format_currency(min_total)} (upper bound TBD)"

        key_parts = []
        if budget_phrase:
            key_parts.append(budget_phrase)
        if duration and str(duration).lower() not in {"tbd", "unknown"}:
            key_parts.append(f"Target window {duration}")
        key_numbers = '; '.join(key_parts) if key_parts else 'Key numbers TBD'

        # Be honest when the call had no renovation scope (e.g. a scheduling call)
        # rather than fabricating a generic "Comprehensive renovation project".
        bullets = [
            f"1. Project Goals – {project_goals}." if project_goals
            else "1. Project Goals – No renovation scope was discussed in this call.",
            f"2. Client Drivers – {client_drivers}." if client_drivers
            else "2. Client Drivers – Not specified.",
            f"3. Key Numbers – {key_numbers}.",
        ]
        for text in bullets:
            para = self.doc.add_paragraph(text, style='List Bullet')
            para.paragraph_format.space_after = Pt(6)

    def _add_property_details(self, data: Dict[str, Any]):
        """Add property details section"""
        self._heading('Property Details', level=1)

        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        table.autofit = True

        property_info = data.get('property_details', {}) or {}

        # Price (fallback to last sold price)
        price_num = _leading_money(property_info.get('price'))
        last_sold_price = property_info.get('last_sold_price')
        last_sold_date = property_info.get('last_sold_date')
        last_sold_num = _leading_money(last_sold_price)

        # Label the row honestly: a genuine active-listing price is "Current
        # Price", but a historical sale (the usual case for researched off-market
        # homes) is "Last Sold Price" — never present a years-old purchase as the
        # current value. Append the sale date when known. Using _leading_money
        # normalizes a model-annotated price ("$736,000 (sold Sep 25...)") to a
        # clean dollar figure so the date isn't rendered twice.
        _has_date = last_sold_date not in (None, '', 'Information not available')
        if price_num is not None:
            price_label, price_display = 'Current Price', f"${price_num:,.2f}"
        elif property_info.get('price') not in (None, '', 'Information not available'):
            price_label, price_display = 'Current Price', str(property_info.get('price'))
        elif last_sold_num is not None:
            price_label = 'Last Sold Price'
            price_display = f"${last_sold_num:,.2f} ({last_sold_date})" if _has_date else f"${last_sold_num:,.2f}"
        elif last_sold_price not in (None, '', 'Information not available'):
            price_label = 'Last Sold Price'
            price_display = f"{last_sold_price} ({last_sold_date})" if _has_date else str(last_sold_price)
        else:
            price_label, price_display = 'Current Price', 'Contact broker for pricing'

        # HOA fee
        hoa_num = _to_number(property_info.get('hoa_fee'))
        hoa_raw = property_info.get('hoa_fee')
        if hoa_num is not None:
            hoa_fee_display = f"${hoa_num:,.2f}/month"
        elif hoa_raw not in (None, '', 'Information not available'):
            hoa_fee_display = str(hoa_raw)
        else:
            hoa_fee_display = 'Information not available'

        def safe_get(info, key):
            return info[key] if key in info and info[key] not in (None, '', []) else 'Information not available'

        bedrooms = safe_get(property_info, 'bedrooms')
        bathrooms = safe_get(property_info, 'bathrooms')
        sqft = property_info.get('sqft')
        sqft_display = f"{sqft} sq ft" if sqft not in (None, '', [], 'Information not available') else 'Information not available'
        year_built = safe_get(property_info, 'year_built')
        property_type = safe_get(property_info, 'property_type')
        if isinstance(property_type, str) and property_type != 'Information not available':
            # Humanize raw enum-style values like "single_family" / "MULTI_FAMILY".
            if '_' in property_type or property_type.isupper() or property_type.islower():
                property_type = property_type.replace('_', ' ').title()
        neighborhood = safe_get(property_info, 'neighborhood')

        details = [
            (price_label, price_display),
            ('Square Footage', sqft_display),
            ('Bedrooms', bedrooms),
            ('Bathrooms', bathrooms),
            ('Year Built', year_built),
            ('Property Type', property_type),
            ('Lot Size', safe_get(property_info, 'lot_size')),
            ('Assessed Value', safe_get(property_info, 'assessed_value')),
            ('Property Taxes', safe_get(property_info, 'property_taxes')),
            ('HOA Fee', hoa_fee_display),
            ('Neighborhood', neighborhood)
        ]

        for item, detail in details:
            row_cells = table.add_row().cells
            row_cells[0].text = item
            row_cells[1].text = str(detail)

    def _download_image(self, url: str) -> Optional[BytesIO]:
        """Download an image and return it as a BytesIO (converted to PNG if needed)."""
        try:
            response = requests.get(url, timeout=IMAGE_DOWNLOAD_TIMEOUT)
            if response.status_code != 200:
                logger.error("Image download failed: %s (status %s)", url, response.status_code)
                return None
            try:
                img = Image.open(BytesIO(response.content))
                if img.format != 'PNG':
                    output = BytesIO()
                    img.convert('RGB').save(output, format='PNG')
                    output.seek(0)
                    return output
                return BytesIO(response.content)
            except Exception as pil_e:
                logger.error("PIL could not process image %s: %s", url, pil_e)
                return BytesIO(response.content)
        except Exception as e:
            logger.error("Exception downloading image %s: %s", url, e)
        return None

    def _add_client_details(self, data: Dict[str, Any]):
        """Add client details section"""
        self._heading('Client Details', level=1)

        transcript_info = data.get('transcript_info', {}) or {}
        client_info = transcript_info.get('client_info', {}) or {}

        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        table.autofit = True

        names = self._lines(client_info.get('names'))
        names_str = ', '.join(names) if names else 'N/A'

        # Preferences (guard against non-dict)
        preferences = client_info.get('preferences', {})
        pref_items = []
        if isinstance(preferences, dict):
            for key, value in preferences.items():
                if value:
                    pref_items.append(f"{key.replace('_', ' ').title()}: {value}")
        preferences_str = '\n'.join(pref_items) if pref_items else 'N/A'

        constraints = self._lines(client_info.get('constraints'))
        constraints_str = '\n'.join(constraints) if constraints else 'N/A'

        # Red flags (guard against non-dict)
        red_flags = client_info.get('red_flags', {})
        flag_items = []
        if isinstance(red_flags, dict):
            for key, value in red_flags.items():
                if value:
                    flag_items.append(key.replace('_', ' ').title())
        red_flags_str = '\n'.join(flag_items) if flag_items else 'None'

        details = [
            ('Name', names_str),
            ('Phone', client_info.get('phone') or 'N/A'),
            ('Email', client_info.get('email') or 'N/A'),
            ('Profession', client_info.get('profession') or 'N/A'),
            ('Preferences', preferences_str),
            ('Constraints', constraints_str),
            ('Potential Concerns', red_flags_str)
        ]

        for item, detail in details:
            row_cells = table.add_row().cells
            row_cells[0].text = item
            row_cells[1].text = str(detail)

    def _add_hyperlink(self, paragraph, text: str, url: str) -> None:
        """Add a hyperlink to a paragraph"""
        part = paragraph._parent.part
        r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        c = OxmlElement('w:color')
        c.set(qn('w:val'), '0000FF')
        rPr.append(c)

        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)

    def _add_property_links(self, data: Dict[str, Any]):
        """Add property links section"""
        self._heading('Property Links', level=1)

        property_details = data.get('property_details') or {}

        realtor_url = data.get('realtor_url')
        if not realtor_url:
            pid = data.get('property_id')
            address = data.get('property_address', '')
            if pid and address:
                parts = [p.strip() for p in address.split(',') if p.strip()]
                if len(parts) >= 3:
                    street = parts[0].replace(' ', '-')
                    city = parts[-2].replace(' ', '-')
                    state_zip = parts[-1].split()
                    if len(state_zip) >= 2:
                        state = state_zip[0]
                        zip_code = state_zip[1]
                        realtor_url = f"https://www.realtor.com/realestateandhomes-detail/{street}_{city}_{state}_{zip_code}_M{pid}"

        link_para = self.doc.add_paragraph()
        link_para.add_run('• Realtor.com: ').bold = True
        if realtor_url:
            self._add_hyperlink(link_para, 'View Listing', realtor_url)
        else:
            link_para.add_run('Not found on Realtor.com')

        # When the property isn't on Realtor, link the Zillow listing as a fallback.
        zillow_url = data.get('zillow_url') or self._zillow_search_url(data.get('property_address'))
        if not realtor_url and zillow_url:
            zpara = self.doc.add_paragraph()
            zpara.add_run('• Zillow: ').bold = True
            self._add_hyperlink(zpara, 'View Listing', zillow_url)

        # Floor Plan section
        self._heading('Floor Plan', level=2)
        floor_plans_data = data.get('floor_plans') or {}
        floor_plans = floor_plans_data.get('floor_plans', []) or property_details.get('floor_plans', [])

        if not floor_plans:
            zillow_url = data.get('zillow_url') or self._zillow_search_url(data.get('property_address'))
            if zillow_url:
                fp = self.doc.add_paragraph()
                fp.add_run('No floor plan from Realtor. Check the Zillow listing (often includes a floor plan): ')
                self._add_hyperlink(fp, 'View on Zillow', zillow_url)
            else:
                self.doc.add_paragraph('No floor plans available.')
            return

        for plan in floor_plans:
            url = plan.get('url') if isinstance(plan, dict) else None
            if not url:
                continue
            image_stream = self._download_image(url)
            if image_stream:
                try:
                    self.doc.add_picture(image_stream, width=Inches(6.0))
                    link_para = self.doc.add_paragraph()
                    link_para.add_run('Floor plan link: ').bold = True
                    self._add_hyperlink(link_para, url, url)
                    self.doc.add_paragraph()
                    continue
                except Exception as e:
                    logger.error("Exception embedding floor-plan image %s: %s", url, e)
            # Could not download/embed — link out instead.
            link_para = self.doc.add_paragraph()
            link_para.add_run('Floor plan link: ').bold = True
            self._add_hyperlink(link_para, url, url)

    def _add_section_break(self):
        """Add a section break for clarity"""
        self.doc.add_paragraph()

    def _finalize_tables(self):
        """Style every table for a clean, branded look that renders reliably.

        Tables are LEFT-anchored at the margin and left to AUTOFIT the page,
        rather than pinned to fixed column widths. Lightweight mobile previews
        (iOS / SharePoint quick-view) render auto-fitted tables correctly but
        mis-place fixed-width tables, pushing them off the page. Earlier reports
        used autofit and previewed correctly, so we keep that and only layer on
        color + borders (which have no layout effect).
        """
        for table in self.doc.tables:
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = True
            _table_borders(table, BORDER_COLOR)
            ncols = len(table.columns)
            for r_idx, row in enumerate(table.rows):
                cells = row.cells
                for cell in cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                if ncols == 2:
                    # Bold + shade the label column; zebra-stripe value cells.
                    for paragraph in cells[0].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                    _shade_cell(cells[0], LABEL_FILL)
                    if r_idx % 2 == 1:
                        _shade_cell(cells[1], ZEBRA_FILL)
                elif ncols == 3:
                    if r_idx == 0:
                        # Header row: light fill + bold DARK text. White-on-dark
                        # vanishes in renderers that drop cell shading (e.g. the
                        # SharePoint preview), so keep the text dark and legible
                        # whether or not the shading renders.
                        for cell in cells:
                            _shade_cell(cell, LABEL_FILL)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                                    run.font.color.rgb = HEADING_COLOR
                    elif r_idx % 2 == 0:
                        for cell in cells:
                            _shade_cell(cell, ZEBRA_FILL)
                    # Right-align the Amount column for tidy figures.
                    for paragraph in cells[1].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _add_building_requirements(self, data: Dict[str, Any]):
        """Add building requirements section from transcript data only (no hard-coded rules)."""
        self._heading('Building Requirements', level=1)

        transcript_info = data.get('transcript_info', {}) or {}
        property_info = transcript_info.get('property_info', {}) or {}

        rules_list = self._lines(property_info.get('building_rules'))
        features_list = self._lines(property_info.get('building_features'))

        if not rules_list and not features_list:
            self.doc.add_paragraph('N/A')
            return

        for rule in rules_list:
            self.doc.add_paragraph(f"• {rule}")
        if features_list:
            self.doc.add_paragraph('Building features: ' + ', '.join(features_list))

    def _add_renovation_scope(self, data: Dict[str, Any]):
        """Add renovation scope section"""
        self._heading('Renovation Scope', level=1)

        transcript_info = data.get('transcript_info') or {}
        renovation = transcript_info.get('renovation_scope') or {}

        rendered_any = False
        for section_name, section in renovation.items():
            if section_name.lower() == 'timeline':
                continue  # handled in dedicated timeline section
            if not isinstance(section, dict) or not self._is_meaningful(section):
                continue
            rows = []
            for key, value in section.items():
                if not self._is_meaningful(value):
                    continue
                pretty = self._stringify(value)
                if not pretty.strip():
                    continue
                rows.append((key.replace('_', ' ').title(), pretty))
            if not rows:
                continue
            rendered_any = True
            self._heading(section_name.replace('_', ' ').title(), level=2)
            table = self.doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            for label, pretty in rows:
                cells = table.add_row().cells
                cells[0].text = label
                cells[1].text = pretty

        if not rendered_any:
            self.doc.add_paragraph('N/A')

    def _add_materials_design(self, data: Dict[str, Any]):
        """Add a Materials & Design section (skipped entirely when empty)."""
        transcript_info = data.get('transcript_info') or {}
        materials = transcript_info.get('materials_and_design') or {}
        if not isinstance(materials, dict) or not self._is_meaningful(materials):
            return

        rows = []
        for key, value in materials.items():
            if not self._is_meaningful(value):
                continue
            pretty = self._stringify(value)
            if pretty.strip():
                rows.append((key.replace('_', ' ').title(), pretty))
        if not rows:
            return

        self._heading('Materials & Design', level=1)
        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        for label, pretty in rows:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = pretty

    def _add_timeline_phasing(self, data: Dict[str, Any]):
        """Add timeline and phasing section dynamically (no default placeholder values)."""
        self._heading('Timeline & Phasing', level=1)

        transcript_info = data.get('transcript_info') or {}
        renovation_scope = transcript_info.get('renovation_scope') or {}
        timeline = renovation_scope.get('timeline') or {}

        rows = []

        # Scalar fields (note: schema uses 'living_arrangements', not 'coordination')
        key_map = {
            'total_duration': 'Total Duration',
            'phasing': 'Phasing',
            'living_arrangements': 'Occupancy',
        }
        for key, label in key_map.items():
            val = timeline.get(key)
            if val not in (None, '', [], {}):
                rows.append((label, str(val)))

        # Timeline constraints (previously dropped)
        constraints = self._lines(timeline.get('constraints'))
        if constraints:
            rows.append(('Constraints', '\n'.join(constraints)))

        # Key dates (previously dropped entirely)
        key_dates = timeline.get('key_dates') or {}
        if isinstance(key_dates, dict):
            date_labels = {
                'survey_completion': 'Survey Completion',
                'walkthrough_scheduled': 'Walkthrough Scheduled',
                'project_start': 'Project Start',
            }
            for key, label in date_labels.items():
                val = key_dates.get(key)
                if val not in (None, '', [], {}):
                    rows.append((label, str(val)))
            milestones = self._lines(key_dates.get('other_milestones'))
            if milestones:
                rows.append(('Other Milestones', '\n'.join(milestones)))

        if not rows:
            self.doc.add_paragraph('N/A')
            return

        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        for label, val in rows:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = val

    def _add_budget_summary(self, data: Dict[str, Any]):
        """Add budget summary section"""
        self._heading('Budget Summary', level=1)

        transcript_info = data.get('transcript_info') or {}
        renovation = transcript_info.get('renovation_scope') or {}
        kitchen = renovation.get('kitchen', {}) or {}
        bathrooms = renovation.get('bathrooms', {}) or {}
        additional = renovation.get('additional_work', {}) or {}
        estimated_costs = additional.get('estimated_costs') or {}

        details = []
        min_total, max_total = 0.0, 0.0
        have_total = False

        # Kitchen
        krange = (kitchen.get('estimated_cost') or {}).get('range') or {}
        kmin = _to_number(krange.get('min'))
        kmax = _to_number(krange.get('max'))
        if kmin:  # treat 0 / None (template default) as "no kitchen budget given"
            details.append(('Kitchen Total', f"{self._format_currency(kmin)} - {self._format_currency(kmax) if kmax else 'TBD'}"))
            min_total += kmin
            max_total += kmax if kmax else kmin
            have_total = True

        # Bathrooms (cap count defensively)
        bcount = self._safe_int(bathrooms.get('count'))
        bcost = _to_number(bathrooms.get('cost_per_bathroom'))
        if bcount and bcost:
            bcount = max(0, min(bcount, 20))
            for i in range(1, bcount + 1):
                details.append((f"Bathroom {i}", f"~{self._format_currency(bcost)}"))
            min_total += bcost * bcount
            max_total += bcost * bcount
            have_total = True

        # Per-sq-ft rate — display only, NOT summed into the total.
        per_sqft = _to_number(estimated_costs.get('per_sqft_cost'))
        if per_sqft:
            details.append(('Cost per Sq Ft', f"{self._format_currency(per_sqft)}/sq ft"))

        # Architect fees
        arch = estimated_costs.get('architect_fees') or {}
        arch_amt = _to_number(arch.get('estimated_amount')) if isinstance(arch, dict) else None
        arch_pct = _to_number(arch.get('percentage')) if isinstance(arch, dict) else None
        if arch_amt:
            details.append(('Architect Fees', self._format_currency(arch_amt)))
            min_total += arch_amt
            max_total += arch_amt
            have_total = True
        elif arch_pct:
            details.append(('Architect Fees', f"{arch_pct:g}%"))

        # Other simple numeric line items
        handled = {'per_sqft_cost', 'total_estimated_range', 'architect_fees', 'additional_fees'}
        if isinstance(estimated_costs, dict):
            for item, cost in estimated_costs.items():
                if item in handled:
                    continue
                c = _to_number(cost)
                if c:
                    details.append((item.replace('_', ' ').title(), self._format_currency(c)))
                    min_total += c
                    max_total += c
                    have_total = True

        # Additional fees (descriptive, not summed)
        for fee in self._lines(estimated_costs.get('additional_fees')):
            details.append(('Additional Fee', fee))

        # Authoritative total range from extraction, if present; else derive from components.
        tr = estimated_costs.get('total_estimated_range') or {}
        tr_min = _to_number(tr.get('min')) if isinstance(tr, dict) else None
        tr_max = _to_number(tr.get('max')) if isinstance(tr, dict) else None
        if tr_min or tr_max:
            lo = tr_min if tr_min is not None else tr_max
            hi = tr_max if tr_max is not None else tr_min
            details.append(('Estimated Total Range', f"{self._format_currency(lo)} - {self._format_currency(hi)}"))
        elif have_total:
            if max_total and max_total != min_total:
                details.append(('Estimated Total Range', f"{self._format_currency(min_total)} - {self._format_currency(max_total)}"))
            elif min_total:
                details.append(('Estimated Total (min)', self._format_currency(min_total)))

        if not details:
            self.doc.add_paragraph('N/A')
            return

        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        for item, detail in details:
            row_cells = table.add_row().cells
            row_cells[0].text = item
            row_cells[1].text = str(detail)

    def _add_project_management(self, data: Dict[str, Any]):
        """Add a Project Management section (skipped entirely when empty)."""
        transcript_info = data.get('transcript_info') or {}
        pm = transcript_info.get('project_management') or {}
        if not isinstance(pm, dict) or not self._is_meaningful(pm):
            return

        rows = []
        for key, value in pm.items():
            # Internal sales-prep brief — skip our own company self-description.
            if key == 'company_details':
                continue
            if not self._is_meaningful(value):
                continue
            pretty = self._stringify(value)
            if pretty.strip():
                rows.append((key.replace('_', ' ').title(), pretty))
        if not rows:
            return

        self._heading('Project Management', level=1)
        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        for label, pretty in rows:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = pretty

    def _add_notes(self, data: Dict[str, Any]):
        """Add notes section"""
        self._heading('Notes:', level=1)

        transcript_info = data.get('transcript_info', {}) or {}
        renovation = transcript_info.get('renovation_scope', {}) or {}
        materials = transcript_info.get('materials_and_design', {}) or {}
        project_mgmt = transcript_info.get('project_management', {}) or {}

        notes = []

        if materials.get('sourcing_responsibility'):
            notes.append(f"• {materials['sourcing_responsibility']}")

        kitchen_reno = renovation.get('kitchen') or {}
        notes.extend(f"• {req}" for req in self._lines(kitchen_reno.get('specific_requirements')))
        bathrooms_reno = renovation.get('bathrooms') or {}
        notes.extend(f"• {req}" for req in self._lines(bathrooms_reno.get('specific_requirements')))

        if project_mgmt.get('communication_preferences'):
            notes.append(f"• Communication preference: {project_mgmt['communication_preferences']}")
        notes.extend(f"• {doc}" for doc in self._lines(project_mgmt.get('documentation_needs')))

        notes = [self._scrub_company(n) for n in notes]
        notes = [n for n in notes if n.strip() and n.lower().strip('• ') not in {'unknown', 'n/a', 'none'}]
        if not notes:
            notes = ['• No special notes recorded.']

        for note in notes:
            self.doc.add_paragraph(note)

    def _add_neighboring_projects(self, data: Dict[str, Any]):
        """Add neighboring projects section from Zoho CRM"""
        self._heading('Neighboring Projects', level=1)

        neighboring_projects = data.get('neighboring_projects', []) or []

        if not neighboring_projects:
            self.doc.add_paragraph("No neighboring projects found in this area.")
            return

        same_building = [p for p in neighboring_projects if p.get('is_same_building')]
        if same_building:
            intro = f"Found {len(same_building)} project(s) in the same building"
            if len(neighboring_projects) > len(same_building):
                intro += f" and {len(neighboring_projects) - len(same_building)} in the neighborhood."
            else:
                intro += "."
        else:
            intro = f"Found {len(neighboring_projects)} project(s) in the neighborhood."

        self.doc.add_paragraph(intro)
        self.doc.add_paragraph()

        # Cap the rendered rows for a clean report (list is pre-sorted by relevance).
        display = neighboring_projects[:MAX_NEIGHBORING_PROJECTS]

        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        header_cells = table.rows[0].cells
        for i, header in enumerate(['Project Address', 'Amount', 'Stage']):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        for project in display:
            row_cells = table.add_row().cells
            row_cells[0].text = str(project.get('deal_name', 'N/A'))

            amount = _to_number(project.get('amount'))
            row_cells[1].text = self._format_currency(amount) if (amount is not None and amount > 0) else "TBD"

            row_cells[2].text = str(project.get('stage', 'Unknown'))

        if len(neighboring_projects) > len(display):
            self.doc.add_paragraph(
                f"Showing the {len(display)} most relevant of {len(neighboring_projects)} projects in this neighborhood."
            )

    def _add_site_feasibility(self, data: Dict[str, Any]):
        """Zoning, flood zone, and renovation-feasibility notes from public-records research."""
        NA = (None, '', 'Information not available')
        zoning = data.get('research_zoning')
        flood = data.get('research_flood')
        notes = [str(n).strip() for n in (data.get('research_feasibility') or []) if str(n).strip()]
        sources = [str(s).strip() for s in (data.get('research_sources') or []) if str(s).strip()]
        has_kv = zoning not in NA or flood not in NA
        if not has_kv and not notes:
            return  # nothing researched — skip the section entirely
        self._heading('Site & Feasibility', level=1)
        if has_kv:
            table = self.doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            table.autofit = True
            for label, val in (('Zoning', zoning), ('Flood Zone', flood)):
                if val not in NA:
                    cells = table.add_row().cells
                    cells[0].text = label
                    cells[1].text = str(val)
        for n in notes:
            self.doc.add_paragraph(f"• {n}")
        if sources:
            # Reduce each source to a display-safe token (domain for URLs) so a
            # long raw URL can't overflow the right margin on the SharePoint/iOS
            # docx preview, which does not break within an unbroken URL token.
            from urllib.parse import urlparse

            def _src_label(s):
                s = str(s).strip()
                if s.startswith(('http://', 'https://')):
                    netloc = (urlparse(s).netloc or s).strip()
                    return netloc[4:] if netloc.startswith('www.') else netloc
                return s if len(s) <= 40 else s[:40].rstrip() + '…'

            labels = []
            for s in sources:
                lbl = _src_label(s)
                if lbl and lbl not in labels:
                    labels.append(lbl)
            p = self.doc.add_paragraph()
            r = p.add_run('Research sources: ' + '; '.join(labels))
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    def _add_owner_profile(self, data: Dict[str, Any]):
        """Ownership + public professional context (no personal/financial/family/social detail)."""
        summary = data.get('owner_summary')
        if not summary or str(summary).strip() in ('', 'Information not available'):
            return
        self._heading('Owner Profile', level=1)
        # The research composes owner_summary as short labeled lines separated by
        # newlines (Ownership / Tenure / Profession / …). Render each as its own
        # bullet; fall back to a single paragraph if it came back as prose.
        lines = [ln.strip().lstrip('•-').strip() for ln in str(summary).splitlines() if ln.strip()]
        if len(lines) > 1:
            for ln in lines:
                self.doc.add_paragraph(f"• {ln}")
        else:
            self.doc.add_paragraph(str(summary).strip())

    def _add_crm_notes(self, data: Dict[str, Any]):
        """Relevant notes pulled from the matching Zoho CRM deal."""
        notes = [str(n).strip() for n in (data.get('zoho_notes') or []) if str(n).strip()]
        if not notes:
            return
        self._heading('CRM Notes', level=1)
        for n in notes:
            self.doc.add_paragraph(f"• {n}")

    def generate_report(self, data: Dict[str, Any], output_dir: str = "data", file_name: str = None) -> Optional[str]:
        """Generate the pre-walkthrough report.

        Each section is rendered in its own try/except so a single bad field can
        never abort the entire report.
        """
        try:
            output_dir = Path(output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)

            sections = [
                ('Header', self._add_header),
                ('Executive Summary', self._add_executive_summary),
                ('Property Details', self._add_property_details),
                ('Site & Feasibility', self._add_site_feasibility),
                ('Client Details', self._add_client_details),
                ('Owner Profile', self._add_owner_profile),
                ('Property Links', self._add_property_links),
                ('Building Requirements', self._add_building_requirements),
                ('Renovation Scope', self._add_renovation_scope),
                ('Materials & Design', self._add_materials_design),
                ('Timeline & Phasing', self._add_timeline_phasing),
                ('Budget Summary', self._add_budget_summary),
                ('Project Management', self._add_project_management),
                ('Neighboring Projects', self._add_neighboring_projects),
                ('CRM Notes', self._add_crm_notes),
                ('Notes', self._add_notes),
            ]

            for name, render in sections:
                try:
                    render(data)
                    self._add_section_break()
                except Exception as e:
                    logger.error("Failed to render report section '%s': %s", name, e, exc_info=True)
                    self.doc.add_paragraph(f"[Section unavailable: {name}]")
                    self._add_section_break()

            # Uniform, symmetric table layout across the whole report.
            self._finalize_tables()

            if not file_name:
                file_name = f"PreWalk_{self._sanitize_filename(str(data.get('property_address', 'report')).split(',')[0])}.docx"
            else:
                # Ensure a caller-supplied name is also safe.
                stem = file_name[:-5] if file_name.lower().endswith('.docx') else file_name
                file_name = f"{self._sanitize_filename(stem)}.docx"

            output_path = output_dir / file_name
            self.doc.save(str(output_path))
            logger.info("Report generated successfully: %s", output_path)
            return str(output_path)

        except Exception as e:
            logger.error("Error generating report: %s", e, exc_info=True)
            return None

    @staticmethod
    def _sanitize_filename(text: str, max_length: int = 100) -> str:
        """Produce a safe, non-empty filename stem from arbitrary text."""
        text = str(text) if text is not None else ''
        invalid_chars = ['/', '\\', ':', '*', '?', '"', '<', '>', '|', '#']
        for char in invalid_chars:
            text = text.replace(char, '_')
        # Strip control characters
        text = ''.join(ch for ch in text if ch.isprintable())
        text = text.replace(' ', '_')
        while '__' in text:
            text = text.replace('__', '_')
        text = text.strip('_. ')
        if len(text) > max_length:
            text = text[:max_length].strip('_. ')
        return text or 'report'

    # ------------------------------------------------------------------
    # Helper to stringify complex values for tables
    # ------------------------------------------------------------------
    def _scrub_company(self, text):
        """Strip references to our own company. This is an INTERNAL brief for the
        Chapter rep attending the walkthrough — not a client proposal — so
        'Chapter will…' / 'Contractor (Chapter) handles…' self-description is noise."""
        if not text:
            return text
        s = str(text)
        for token in ('(Chapter)', '(chapter)', 'Chapter Renovations', 'Chapter'):
            s = s.replace(token, '')
        while '  ' in s:
            s = s.replace('  ', ' ')
        return s.strip().strip(';,').strip()

    def _format_range(self, d):
        """Format a {min,max}/{low,high} numeric dict as a currency range, else None."""
        if not isinstance(d, dict) or not d:
            return None
        keys = {str(k).lower() for k in d.keys()}
        if not keys <= {'min', 'max', 'low', 'high'}:
            return None
        lo = _to_number(d.get('min', d.get('low')))
        hi = _to_number(d.get('max', d.get('high')))
        if lo is None and hi is None:
            return None
        if lo is not None and hi is not None:
            return self._format_currency(lo) if lo == hi else f"{self._format_currency(lo)} – {self._format_currency(hi)}"
        return f"From {self._format_currency(lo)}" if lo is not None else f"Up to {self._format_currency(hi)}"

    def _stringify(self, value):
        """Human-readable string for a value, with our own company name scrubbed."""
        return self._scrub_company(self._stringify_raw(value))

    def _stringify_raw(self, value):
        """Convert list/dict to a human-readable string, dropping empty/zero parts."""
        if isinstance(value, (list, tuple)):
            return ', '.join(str(v) for v in value if self._is_meaningful(v))
        if isinstance(value, dict):
            parts = []
            for k, v in value.items():
                if not self._is_meaningful(v):
                    continue
                label = str(k).replace('_', ' ').title()
                if isinstance(v, dict):
                    rng = self._format_range(v)
                    if rng is not None:
                        parts.append(f"{label}: {rng}")
                        continue
                    sub = ', '.join(
                        f"{str(sk).replace('_', ' ')}: {sv}"
                        for sk, sv in v.items() if self._is_meaningful(sv)
                    )
                    if sub:
                        parts.append(f"{label} → {sub}")
                elif isinstance(v, (list, tuple)):
                    items = ', '.join(str(i) for i in v if self._is_meaningful(i))
                    if items:
                        parts.append(f"{label}: {items}")
                else:
                    parts.append(f"{label}: {v}")
            return '; '.join(parts)
        return str(value)
